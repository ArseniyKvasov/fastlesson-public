from datetime import timedelta

from celery.utils.time import remaining
from django.core.exceptions import PermissionDenied
from django.utils import timezone

from aiogram import Router, types, F
from aiogram.exceptions import TelegramBadRequest
from aiogram.filters import Command
from aiogram.types import InlineKeyboardButton, InlineKeyboardMarkup
from aiogram.utils.keyboard import InlineKeyboardBuilder
from aiogram.fsm.context import FSMContext
from aiogram.fsm.state import StatesGroup, State
from asgiref.sync import sync_to_async
from typing import Tuple
import tempfile
import os
import re
import html
from typing import Optional

import markdown
from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from fastlesson_bot.services.rate_limit import check_rate_limit
from fastlesson_bot.services.user_service import (
    get_user_by_tg,
    set_user_subject,
    set_user_level,
    create_lesson_for_user, track_user_activity, can_generate_lesson, get_or_create_user,
)
from core.models import SubjectChoices, LevelChoices, Lesson, GenerationStatus, ImproveStatus, UserRole
from metrics.models import SupportTicket, TicketMessage
from core.tasks import generate_lesson_task, improve_block_task

router = Router()


# -----------------------------
# FSM States
# -----------------------------
class CreateLessonStates(StatesGroup):
    waiting_for_title = State()
    waiting_for_new_title = State()  # для изменения темы существующего урока

_FRAC_LATEX_RE = re.compile(r'\\frac\s*\{\s*([^{}]+?)\s*\}\s*\{\s*([^{}]+?)\s*\}')
_SQRT_LATEX_RE = re.compile(r'\\sqrt\s*\{\s*([^{}]+?)\s*\}')
_TIMES_LATEX_RE = re.compile(r'\\times|\\cdot')
_NOT_EQUAL_RE = re.compile(r'\\neq')
_SUP_RE = re.compile(r'\^\{([^{}]+?)\}')
_SUB_RE = re.compile(r'_\{([^{}]+?)\s*\}')
_INLINE_DOLLAR_RE = re.compile(r'\$(.*?)\$', flags=re.S)
_BLOCK_DOLLAR_RE = re.compile(r'\$\$(.*?)\$\$', flags=re.S)

ALLOWED_TAGS = {"b", "strong", "i", "em", "br", "p", "ul", "ol", "li", "table", "tr", "td", "th", "tbody", "thead", "h1", "h2", "h3", "h4", "h5", "h6", "math"}


def _convert_latex_to_text(s: str) -> str:
    if not s:
        return ""
    s = re.sub(r"\\(?:left|right)\b", "", s)
    s = _NOT_EQUAL_RE.sub("≠", s)

    def _frac_repl(m: re.Match) -> str:
        num = m.group(1).strip()
        den = m.group(2).strip()
        # добавляем скобки, если числитель содержит +, -, *, /
        if re.search(r'[+\-\*/\/]', num):
            num = f"({num})"
        return f"{num}/{den}"

    for _ in range(10):
        new = _FRAC_LATEX_RE.sub(_frac_repl, s)
        if new == s:
            break
        s = new

    s = _SQRT_LATEX_RE.sub(lambda m: f"√({m.group(1).strip()})", s)
    s = _TIMES_LATEX_RE.sub("×", s)
    s = _SUP_RE.sub(lambda m: f"^{m.group(1).strip()}", s)
    s = _SUB_RE.sub(lambda m: f"_{m.group(1).strip()}", s)
    s = re.sub(r"\\[a-zA-Z]+", "", s)

    def _collapse_spaces(match: re.Match) -> str:
        text = match.group(0)
        return " " if "\n" not in text else text

    s = re.sub(r"[ \t\r\f\v]+", _collapse_spaces, s)
    lines = [ln.strip() for ln in s.split('\n')]
    s = '\n'.join(lines)
    return s.strip()


def sanitize_math_to_text_fragment(s: Optional[str]) -> str:
    if not s:
        return ""
    t = s
    t = _BLOCK_DOLLAR_RE.sub(lambda m: _convert_latex_to_text(m.group(1)), t)
    t = _INLINE_DOLLAR_RE.sub(lambda m: _convert_latex_to_text(m.group(1)), t)
    t = _convert_latex_to_text(t)
    return t


def sanitize_html(text: Optional[str]) -> str:
    if not text:
        return ""
    text = html.unescape(text)
    text = re.sub(r"(?i)<br\s*/?>", "<br/>", text)
    text = re.sub(r"(?i)</p\s*>", "</p>", text)
    text = re.sub(r"(?i)<p\b[^>]*>", "<p>", text)
    text = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", text, flags=re.S)
    text = re.sub(r"(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)", r"<i>\1</i>", text, flags=re.S)
    def _quote_repl(m: re.Match) -> str:
        inner = m.group(1)
        return f'"<i>{inner}</i>"'
    text = re.sub(r"(?<!\w)'(.*?)'(?!\w)", _quote_repl, text, flags=re.S)
    text = re.sub(r"(?<!\w)`(.*?)`(?!\w)", _quote_repl, text, flags=re.S)
    text = _BLOCK_DOLLAR_RE.sub(lambda m: _convert_latex_to_text(m.group(1)), text)
    text = _INLINE_DOLLAR_RE.sub(lambda m: _convert_latex_to_text(m.group(1)), text)
    text = _convert_latex_to_text(text)
    soup = BeautifulSoup(text, "html.parser")
    for tag in list(soup.find_all()):
        name = tag.name.lower() if tag.name else ""
        if name in ALLOWED_TAGS:
            tag.attrs = {}
        else:
            tag.unwrap()
    return str(soup)

_SUP_MAP = {
    "0": "⁰", "1": "¹", "2": "²", "3": "³", "4": "⁴", "5": "⁵", "6": "⁶",
    "7": "⁷", "8": "⁸", "9": "⁹", "+": "⁺", "-": "⁻", "=": "⁼", "(": "⁽",
    ")": "⁾", "n": "ⁿ",
}

_SUB_MAP = {
    "0": "₀", "1": "₁", "2": "₂", "3": "₃", "4": "₄", "5": "₅", "6": "₆",
    "7": "₇", "8": "₈", "9": "₉", "+": "₊", "-": "₋", "=": "₌", "(": "₍",
    ")": "₎", "a": "ₐ", "e": "ₑ", "o": "ₒ", "x": "ₓ", "i": "ᵢ", "r": "ᵣ",
    "u": "ᵤ", "v": "ᵥ", "t": "ₜ", "n": "ₙ", "h": "ₕ", "k": "ₖ", "l": "ₗ",
    "m": "ₘ", "s": "ₛ", "p": "ₚ", "y": "ᵧ",
}


def _replace_simple_superscripts(text: str) -> str:
    def repl(m):
        chars = m.group(1)
        return ''.join(_SUP_MAP.get(ch, f'^{ch}') for ch in chars)
    text = re.sub(r'\^([0-9n+\-=\(\)])', repl, text)
    return text

def _replace_simple_subscripts(text: str) -> str:
    def repl(m):
        chars = m.group(1)
        return ''.join(_SUB_MAP.get(ch, f'_{ch}') for ch in chars)
    text = re.sub(r'_([0-9a-z+\-=\(\)])', repl, text)
    return text

def _append_runs_from_fragment(paragraph, fragment_html: str):
    frag_soup = BeautifulSoup(fragment_html, "html.parser")

    # Преобразуем LaTeX внутри всех текстовых узлов
    for text_node in frag_soup.find_all(string=True):
        if isinstance(text_node, NavigableString):
            new_text = sanitize_math_to_text_fragment(str(text_node))
            new_text = _replace_simple_superscripts(new_text)
            new_text = _replace_simple_subscripts(new_text)
            if new_text != str(text_node):
                text_node.replace_with(new_text)

    FORMATTING_TAGS = ("b", "strong", "i", "em")
    PUNCTUATION_AFTER = {',', '.', ':', '"', "'", ';', '*', ')', ']', '}', '?', '!', '—', '–'}

    def _first_non_space_char_after(node):
        cur = node
        nxt = cur.next_sibling
        while True:
            if nxt is None:
                cur = cur.parent
                if cur is None:
                    return None
                nxt = cur.next_sibling
                continue

            if isinstance(nxt, NavigableString):
                txt = str(nxt)
                for ch in txt:
                    if ch == '\n':
                        return '\n'
                    if not ch.isspace():
                        return ch
                nxt = nxt.next_sibling
                continue

            if isinstance(nxt, Tag):
                inner = nxt
                while True:
                    children = list(inner.contents)
                    if not children:
                        break
                    first = children[0]
                    if isinstance(first, NavigableString):
                        txt = str(first)
                        for ch in txt:
                            if ch == '\n':
                                return '\n'
                            if not ch.isspace():
                                return ch
                        break
                    elif isinstance(first, Tag):
                        inner = first
                        continue
                    else:
                        break
                nxt = nxt.next_sibling
                continue

            nxt = nxt.next_sibling

    def _walk(node, bold=False, italic=False, prev_tag=None):
        if isinstance(node, NavigableString):
            text = _replace_simple_superscripts(str(node))
            text = _replace_simple_subscripts(text)
            if not text:
                return
            parts = re.split(r'(\s+)', text)
            for part in parts:
                if part == "":
                    continue
                run = paragraph.add_run(part)
                run.bold = bold
                run.italic = italic
            return

        if isinstance(node, Tag):
            tag_name = node.name.lower()
            new_bold = bold or tag_name in FORMATTING_TAGS
            new_italic = italic or tag_name in FORMATTING_TAGS

            if tag_name in FORMATTING_TAGS and prev_tag not in FORMATTING_TAGS:
                paragraph.add_run(' ')

            for child in node.children:
                _walk(child, new_bold, new_italic, prev_tag=tag_name)

            if tag_name in FORMATTING_TAGS:
                next_ch = _first_non_space_char_after(node)
                if next_ch is not None and next_ch != '\n' and next_ch not in PUNCTUATION_AFTER:
                    paragraph.add_run(' ')

    for child in frag_soup.contents:
        _walk(child)


def _add_paragraph(container, html_fragment: str, style: Optional[str] = None):
    p = container.add_paragraph(style=style) if style else container.add_paragraph()
    _append_runs_from_fragment(p, html_fragment)
    return p


def sanitize_word(html_text: Optional[str], doc: Document):
    if not html_text:
        return
    safe_html = sanitize_html(html_text)
    soup = BeautifulSoup(safe_html, "html.parser")
    for node in soup.contents:
        if isinstance(node, NavigableString):
            text = str(node)
            if text:
                _add_paragraph(doc, text)
            continue
        if not isinstance(node, Tag):
            continue
        tag = node.name.lower()
        if tag in ("h1", "h2", "h3", "h4", "h5", "h6"):
            level = int(tag[1]) if len(tag) > 1 and tag[1].isdigit() else 1
            level = min(max(level, 1), 4)
            heading_para = doc.add_heading(level=level)
            _append_runs_from_fragment(heading_para, node.decode_contents())
            continue
        if tag == "p":
            _add_paragraph(doc, node.decode_contents())
            continue
        if tag in ("ul", "ol"):
            is_ordered = tag == "ol"
            li_nodes = node.find_all("li", recursive=False)
            for idx, li in enumerate(li_nodes, start=1):
                li_html = li.decode_contents()
                prefix = f"{idx}. " if is_ordered else "• "
                p = doc.add_paragraph()
                p.add_run(prefix)
                _append_runs_from_fragment(p, li_html)
            continue
        if tag == "table":
            rows = node.find_all("tr")
            if not rows:
                continue
            col_count = max(1, len(rows[0].find_all(["td", "th"])))
            table = doc.add_table(rows=len(rows), cols=col_count)
            table.style = "Table Grid"
            for r, row in enumerate(rows):
                cells = row.find_all(["td", "th"])
                for c in range(col_count):
                    cell = table.cell(r, c)
                    try:
                        cell.text = ""
                    except Exception:
                        pass
                    # Добавляем небольшой внутренний отступ через XML
                    tcPr = cell._tc.get_or_add_tcPr()
                    tcMar = OxmlElement('w:tcMar')
                    for side in ('top', 'start', 'bottom', 'end'):
                        elem = OxmlElement(f'w:{side}')
                        elem.set(qn('w:w'), "100")  # примерно 0.1 дюйма ~ 2.5 мм
                        elem.set(qn('w:type'), "dxa")
                        tcMar.append(elem)
                    tcPr.append(tcMar)

                    if c < len(cells):
                        td_node = cells[c]
                        ps = td_node.find_all("p", recursive=False)
                        if ps:
                            for pnode in ps:
                                _add_paragraph(cell, pnode.decode_contents())
                        else:
                            _add_paragraph(cell, td_node.decode_contents())
            continue
        if tag == "math":
            math_text = node.get_text(" ", strip=False)
            doc.add_paragraph(sanitize_math_to_text_fragment(math_text if math_text else "[formula]"))
            continue
        _add_paragraph(doc, node.decode_contents())



def markdown_to_html(md_text: str) -> str:
    return markdown.markdown(md_text, extensions=["tables"]);




def subject_kb() -> InlineKeyboardBuilder:
    kb = InlineKeyboardBuilder()
    for key, label in SubjectChoices.choices:
        kb.button(text=label, callback_data=f"subject:{key}")
    kb.adjust(2)
    return kb


def level_kb() -> InlineKeyboardBuilder:
    kb = InlineKeyboardBuilder()
    for key, label in LevelChoices.choices:
        kb.button(text=label, callback_data=f"level:{key}")
    kb.adjust(2)
    return kb


def lesson_actions_kb(lesson_id: str, status: Optional[GenerationStatus] = None) -> InlineKeyboardBuilder:
    """Клавиатура действий для детального вида урока.
    Вместо одной кнопки "Действия" показываем варианты "Посмотреть", "Отправить" и "Скачать".
    """
    kb = InlineKeyboardBuilder()
    if not status or status.total == 0:
        kb.button(text="🚀 Сгенерировать", callback_data=f"lesson_generate:{lesson_id}")
        kb.button(text="✏️ Изменить тему", callback_data=f"lesson_change_title:{lesson_id}")
        kb.button(text="❌ Удалить урок", callback_data=f"lesson_delete:{lesson_id}")
        kb.button(text="🏠 На главную", callback_data="main_menu")
    elif status.completed < status.total:
        kb.button(text="🔄 Обновить статус", callback_data=f"lesson_status:{lesson_id}")
        kb.button(text="🏠 На главную", callback_data="main_menu")
    else:
        # Генерация завершена — показываем все действия
        kb.button(text="👀 Посмотреть", callback_data=f"lesson_view:{lesson_id}:1")
        #kb.button(text="📤 Отправить", callback_data=f"lesson_send:{lesson_id}")
        kb.button(text="📥 Скачать", callback_data=f"lesson_download:{lesson_id}")
        kb.button(text="✏️ Изменить тему", callback_data=f"lesson_change_title:{lesson_id}")
        kb.button(text="❌ Удалить урок", callback_data=f"lesson_delete:{lesson_id}")
        kb.button(text="🏠 На главную", callback_data="main_menu")
    kb.adjust(2)
    return kb

class EditBlockStates(StatesGroup):
    waiting_for_new_content = State()


def navigation_kb_for_block(lesson_id: str, block_index: int, total_blocks: int) -> InlineKeyboardBuilder:
    kb = InlineKeyboardBuilder()
    if block_index > 1:
        kb.button(text="⬅️ Назад", callback_data=f"lesson_view:{lesson_id}:{block_index - 1}")
    else:
        kb.button(text="ℹ️ Об уроке", callback_data=f"lesson_detail:{lesson_id}")
    if block_index < total_blocks:
        kb.button(text="➡️ Далее", callback_data=f"lesson_view:{lesson_id}:{block_index + 1}")
    else:
        kb.button(text="ℹ️ Об уроке", callback_data=f"lesson_detail:{lesson_id}")

    # вместо "Удалить" — кнопка "Действия"
    kb.button(text="⚙️ Действия", callback_data=f"lesson_actions:{lesson_id}:{block_index}")
    #kb.button(text="📤 Отправить", callback_data=f"lesson_send:{lesson_id}")
    kb.button(text="📥 Скачать", callback_data=f"lesson_download:{lesson_id}")
    kb.button(text="🏠 На главную", callback_data="main_menu")
    kb.adjust(2)
    return kb

# -----------------------
# Клавиатура действий для конкретного блока
# -----------------------
def actions_kb_for_block(lesson_id: str, block_index: int) -> InlineKeyboardBuilder:
    kb = InlineKeyboardBuilder()
    kb.button(text="✏️ Редактировать", callback_data=f"lesson_edit:{lesson_id}:{block_index}")
    kb.button(text="🤖 Улучшить ИИ", callback_data=f"lesson_ai_menu:{lesson_id}:{block_index}")
    kb.button(text="❌ Удалить", callback_data=f"lesson_delete_block:{lesson_id}:{block_index}")
    kb.button(text="◀️ Назад к блоку", callback_data=f"lesson_view:{lesson_id}:{block_index}")
    kb.adjust(2)
    return kb

# -----------------------------
# DB-safe wrappers
# -----------------------------
async def safe_get_lesson(lesson_id: str) -> Optional[Lesson]:
    try:
        return await sync_to_async(Lesson.objects.get)(id=lesson_id)
    except Lesson.DoesNotExist:
        return None


async def safe_get_status(lesson: Lesson) -> Optional[GenerationStatus]:
    try:
        return await sync_to_async(GenerationStatus.objects.get)(lesson=lesson)
    except GenerationStatus.DoesNotExist:
        return None


def format_subject_level_labels(subject_key: Optional[str], level_key: Optional[str]) -> Tuple[str, str]:
    subject_label = SubjectChoices(subject_key).label if subject_key else "—"
    level_label = LevelChoices(level_key).label if level_key else "—"
    return subject_label, level_label


async def safe_edit_text(message: types.Message, text: str, **kwargs) -> None:
    try:
        await message.edit_text(text, **kwargs)
    except TelegramBadRequest as e:
        # мелкие ошибки игнорируем, большие — пробрасываем
        if "message is not modified" in str(e):
            return
        raise


# -----------------------------
# Handlers: Subject & Level selection flow
# -----------------------------
@router.callback_query(F.data == "choose_subject")
async def choose_subject(callback: types.CallbackQuery):
    try:
        await callback.message.edit_reply_markup(reply_markup=None)
    except Exception:
        pass

    await safe_edit_text(
        callback.message,
        "Выберите предмет, который вы преподаёте:",
        reply_markup=subject_kb().as_markup()
    )

    await callback.answer()


@router.callback_query(F.data.startswith("subject:"))
async def subject_selected(callback: types.CallbackQuery):
    subject_key = callback.data.split(":", 1)[1]
    user = await get_or_create_user(callback.from_user.id, UserRole.SCHOOL_TEACHER, telegram_username=callback.from_user.username)
    await set_user_subject(callback.from_user.id, subject_key)

    label = sanitize_html(SubjectChoices(subject_key).label)

    # Обновляем текст старого сообщения, убираем старые подсказки
    try:
        await safe_edit_text(
            callback.message,
            f"Предмет сохранён: <b>{label}</b>",
            parse_mode="HTML",
            reply_markup=None
        )
    except Exception:
        pass  # если редактировать не получилось — ничего страшного

    # Отправляем сразу клавиатуру для выбора уровня
    await callback.message.answer(
        "Выберите уровень учеников:",
        reply_markup=level_kb().as_markup()
    )

    await callback.answer()


@router.callback_query(F.data.startswith("level:"))
async def level_selected(callback: types.CallbackQuery, state: FSMContext):
    level_key = callback.data.split(":", 1)[1]
    user = await get_or_create_user(callback.from_user.id, UserRole.SCHOOL_TEACHER, telegram_username=callback.from_user.username)
    await set_user_level(callback.from_user.id, level_key)

    label = sanitize_html(LevelChoices(level_key).label)

    # Редактируем сообщение с предметом, чтобы убрать "Теперь выберите уровень учеников"
    try:
        await safe_edit_text(
            callback.message,
            f"Выбран уровень учеников: <b>{label}</b>\n\nВведите тему урока:",
            parse_mode="HTML"
        )
    except Exception:
        # если сообщение не редактируется (например, текст совпадает), ничего не делаем
        pass

    await state.set_state(CreateLessonStates.waiting_for_title)

    await callback.answer()


@router.message(CreateLessonStates.waiting_for_title)
async def receive_lesson_title(message: types.Message, state: FSMContext):
    title = message.text.strip()
    if not title:
        await message.answer("Пожалуйста, введите тему урока (не пустую строку).")
        return

    try:
        await get_or_create_user(message.from_user.id, UserRole.SCHOOL_TEACHER, telegram_username=message.from_user.username)
    except Exception as e:
        pass

    user = await get_user_by_tg(message.from_user.id)
    metrics = await sync_to_async(track_user_activity)(user)
    subject_label, level_label = format_subject_level_labels(user.subject, user.level)

    lesson = await create_lesson_for_user(
        message.from_user.id,
        title,
        user.subject,
        user.level
    )

    kb = lesson_actions_kb(str(lesson.id))
    kb.adjust(2)

    # Отправляем итоговое сообщение и очищаем состояние
    await message.answer(
        f"✅ Урок «{sanitize_html(lesson.title)}» создан!\n\n"
        f"Предмет: <b>{sanitize_html(subject_label)}</b>\n"
        f"Уровень: <b>{sanitize_html(level_label)}</b>\n\n"
        f"Выберите действие:",
        parse_mode="HTML",
        reply_markup=kb.as_markup()
    )

    await state.clear()



@router.callback_query(F.data.startswith("lesson_change_title:"))
async def prompt_change_title(callback: types.CallbackQuery, state: FSMContext):
    lesson_id = callback.data.split(":", 1)[1]
    lesson = await safe_get_lesson(lesson_id)
    if not lesson:
        await callback.answer("⚠️ Урок не найден", show_alert=True)
        return

    try:
        await get_or_create_user(callback.from_user.id, UserRole.SCHOOL_TEACHER, telegram_username=callback.from_user.username)
    except Exception as e:
        pass

    try:
        await callback.message.edit_reply_markup(reply_markup=None)
    except Exception:
        pass

    await safe_edit_text(
        callback.message,
        f"Текущее название: <b>{sanitize_html(lesson.title)}</b>\n\nВведите новое название:",
        parse_mode="HTML",
    )
    await state.update_data(edit_lesson_id=str(lesson.id))
    await state.set_state(CreateLessonStates.waiting_for_new_title)
    await callback.answer()

@router.message(CreateLessonStates.waiting_for_new_title)
async def receive_new_lesson_title(message: types.Message, state: FSMContext):
    new_title = message.text.strip()
    if not new_title:
        await message.answer(
            "❌ Название не может быть пустым. Пожалуйста, введите новое название урока."
        )
        return

    try:
        await get_or_create_user(message.from_user.id, UserRole.SCHOOL_TEACHER, telegram_username=message.from_user.username)
    except Exception as e:
        pass

    # Получаем id урока из FSMContext
    data = await state.get_data()
    lesson_id = data.get("edit_lesson_id")
    if not lesson_id:
        await message.answer("⚠️ Не удалось определить урок для изменения.")
        await state.clear()
        return

    lesson = await safe_get_lesson(lesson_id)
    if not lesson:
        await message.answer("⚠️ Урок не найден.")
        await state.clear()
        return

    # Обновляем название
    lesson.title = new_title
    await sync_to_async(lesson.save)()

    try:
        await message.edit_reply_markup(reply_markup=None)
    except Exception:
        pass

    # Кнопка "Назад к уроку"
    kb = InlineKeyboardBuilder()
    kb.button(text="🔙 Об уроке", callback_data=f"lesson_detail:{lesson.id}")
    kb.adjust(1)

    # Отправляем новое сообщение с обновленным названием
    await message.answer(
        f"✅ Название урока успешно изменено!\n\nНовое название: <b>{sanitize_html(new_title)}</b>",
        parse_mode="HTML",
        reply_markup=kb.as_markup()
    )

    # Очищаем состояние
    await state.clear()



# -----------------------------
# Handlers: Lesson detail, generation and status
# -----------------------------
@router.callback_query(F.data.startswith("lesson_detail:"))
async def lesson_detail(callback: types.CallbackQuery, state: FSMContext):
    await state.clear()
    lesson_id = callback.data.split(":", 1)[1]
    lesson = await safe_get_lesson(lesson_id)
    user = await get_user_by_tg(callback.from_user.id)
    metrics = await sync_to_async(track_user_activity)(user)
    if not lesson:
        await callback.answer("⚠️ Урок не найден", show_alert=True)
        return

    try:
        await get_or_create_user(callback.from_user.id, UserRole.SCHOOL_TEACHER, telegram_username=callback.from_user.username)
    except Exception as e:
        pass

    status = await safe_get_status(lesson)
    subject_label, level_label = format_subject_level_labels(lesson.subject, lesson.level)

    text = (
        f"📘 <b>Урок:</b> {sanitize_html(lesson.title)}\n\n"
        f"Предмет: <b>{sanitize_html(subject_label)}</b>\n"
        f"Уровень: <b>{sanitize_html(level_label)}</b>\n"
    )

    if status and status.total > 0 and status.completed < status.total:
        text += f"\n⏳ Генерация идёт... {status.completed}/{status.total}"

    try:
        await callback.message.edit_reply_markup(reply_markup=None)
    except Exception:
        pass

    kb = lesson_actions_kb(str(lesson.id), status)

    # Новое сообщение вместо редактирования
    await callback.message.answer(
        text,
        parse_mode="HTML",
        reply_markup=kb.as_markup()
    )
    await callback.answer()


@router.callback_query(F.data.startswith("lesson_generate:"))
async def lesson_generate(callback: types.CallbackQuery):
    try:
        await sync_to_async(check_rate_limit)(
            callback.from_user.id,
            "lesson_generate",
            limit=1,
            window=30
        )
    except PermissionDenied as e:
        await callback.answer(f"⚠️ {str(e)}")
        return

    try:
        await get_or_create_user(callback.from_user.id, UserRole.SCHOOL_TEACHER, telegram_username=callback.from_user.username)
    except Exception as e:
        pass

    lesson_id = callback.data.split(":", 1)[1]
    lesson = await safe_get_lesson(lesson_id)
    user = await get_user_by_tg(callback.from_user.id)

    if not await sync_to_async(can_generate_lesson)(callback.from_user.id, lesson):
        await callback.answer("⚠️ Недостаточно прав", show_alert=True)
        return

    # проверяем, остались ли генерации
    remaining = await sync_to_async(lambda: user.remaining_generations)()

    if remaining <= 0:
        kb = InlineKeyboardBuilder()
        kb.button(text="💎 Пополнить генерации", callback_data="shop")
        kb.adjust(1)

        try:
            await callback.message.edit_reply_markup(reply_markup=None)
        except Exception:
            pass

        await safe_edit_text(
            callback.message,
            "⚠️ У вас закончились генерации.\n\n"
            "Пополните генерации <b>и продолжайте создавать задания!</b>",
            parse_mode="HTML",
            reply_markup=kb.as_markup()
        )
        await callback.answer()
        return

    # списываем одну генерацию
    await sync_to_async(user.decrement_generation)()

    # инициализируем метрики
    metrics = await sync_to_async(track_user_activity)(user)
    await sync_to_async(metrics.update_last_generated)()

    if not lesson:
        await callback.answer("⚠️ Урок не найден", show_alert=True)
        return

    # создаём статус
    await sync_to_async(GenerationStatus.objects.create)(
        lesson=lesson,
        total=0,
        completed=0,
    )

    # генерация успешно запущена → теперь можно убрать старые кнопки
    try:
        await callback.message.edit_reply_markup(reply_markup=None)
    except Exception:
        pass

    generate_lesson_task.delay(str(lesson.id))

    kb = InlineKeyboardBuilder()
    kb.button(text="🔄 Обновить статус", callback_data=f"lesson_status:{lesson.id}")
    kb.button(text="🏠 На главную", callback_data="main_menu")
    kb.adjust(1)

    await safe_edit_text(
        callback.message,
        f"⏳ Генерация урока «{sanitize_html(lesson.title)}» началась...",
        parse_mode="HTML",
        reply_markup=kb.as_markup()
    )
    await callback.answer()



@router.callback_query(F.data.startswith("lesson_status:"))
async def check_lesson_status(callback: types.CallbackQuery):
    try:
        await sync_to_async(check_rate_limit)(
            callback.from_user.id,
            "update_status",
            limit=25,
            window=30
        )
    except PermissionDenied as e:
        await callback.answer(f"⚠️ {str(e)}")
        return

    lesson_id = callback.data.split(":", 1)[1]
    try:
        status = await sync_to_async(GenerationStatus.objects.get)(lesson_id=lesson_id)
    except GenerationStatus.DoesNotExist:
        await callback.answer("⚠️ Статус не найден", show_alert=True)
        return

    if status.total == 0 and status.completed == 0:
        text = "⏳ Начинаем генерацию. Это займет меньше двух минут."
    elif status.completed < status.total:
        text = f"⏳ Генерация идёт...\nГотово: {status.completed}/{status.total}"
    else:
        text = f"✅ Генерация завершена!"

    try:
        await callback.message.edit_reply_markup(reply_markup=None)
    except Exception:
        pass

    kb = InlineKeyboardBuilder()
    if status.completed == status.total and status.total > 0:
        kb.button(text="👀 Посмотреть", callback_data=f"lesson_view:{lesson_id}:1")
    else:
        kb.button(text="🔄 Обновить статус", callback_data=f"lesson_status:{lesson_id}")
    kb.button(text="На главную", callback_data="main_menu")
    kb.adjust(1)

    try:
        await safe_edit_text(callback.message, text, reply_markup=kb.as_markup())
    except TelegramBadRequest as e:
        # safe_edit_text уже обрабатывает "message is not modified"
        raise

    await callback.answer()


# -----------------------------
# Unified Lesson View
# -----------------------------
@router.callback_query(F.data.startswith("lesson_view"))
async def lesson_view(callback: types.CallbackQuery):
    parts = callback.data.split(":")
    lesson_id = parts[1] if len(parts) > 1 else None
    block_index = int(parts[2]) if len(parts) > 2 else 1

    try:
        await get_or_create_user(callback.from_user.id, UserRole.SCHOOL_TEACHER, telegram_username=callback.from_user.username)
    except Exception as e:
        pass

    if not lesson_id:
        await callback.answer("⚠️ Неверные данные", show_alert=True)
        return

    lesson = await safe_get_lesson(lesson_id)
    if not lesson:
        await callback.answer("⚠️ Урок не найден", show_alert=True)
        return

    if not await sync_to_async(can_generate_lesson)(callback.from_user.id, lesson):
        await callback.answer("⚠️ Недостаточно прав", show_alert=True)
        return

    try:
        await sync_to_async(mark_lesson_discovered)(lesson)
    except Exception as e:
        print(e)
        pass

    blocks = await sync_to_async(list)(lesson.blocks.order_by("order").all())
    total_blocks = len(blocks)
    if total_blocks == 0:
        await callback.answer("⚠️ Урок пустой", show_alert=True)
        return

    block_index = max(1, min(block_index, total_blocks))
    block = blocks[block_index - 1]

    text = (
        f"📘 <b>Урок:</b> {sanitize_html(lesson.title)[:100]}\n\n"
        f"<b>Блок {block.order}</b> — {sanitize_html(block.title)[:100]}\n\n"
        f"{sanitize_html(block.content.replace('#', ''))[:3800]}\n\n"
        f"<i>Блок {block_index}/{total_blocks}</i>"
    )

    try:
        await callback.message.edit_reply_markup(reply_markup=None)
    except Exception:
        pass

    kb = navigation_kb_for_block(str(lesson.id), block_index, total_blocks)

    try:
        await safe_edit_text(
            callback.message,
            text,
            parse_mode="HTML",
            reply_markup=kb.as_markup()
        )
    except TelegramBadRequest as e:
        raise

    await callback.answer()


# -----------------------------
# Placeholders for other actions (delete, send, download)
# -----------------------------
# 1️⃣ Шаг: Запрос подтверждения удаления
@router.callback_query(F.data.startswith("lesson_delete:"))
async def lesson_delete_confirm(callback: types.CallbackQuery):
    lesson_id = callback.data.split(":", 1)[1]
    lesson = await safe_get_lesson(lesson_id)
    if not lesson:
        await callback.answer("⚠️ Урок не найден", show_alert=True)
        return

    try:
        await callback.message.edit_reply_markup(reply_markup=None)
    except Exception:
        pass

    # Клавиатура подтверждения
    kb = InlineKeyboardBuilder()
    kb.button(text="✅ Удалить", callback_data=f"lesson_delete_confirm:{lesson_id}")
    kb.button(text="❌ Отмена", callback_data=f"lesson_detail:{lesson_id}")
    kb.adjust(2)

    # Отправляем сообщение с подтверждением
    await callback.message.answer(
        f"⚠️ Вы действительно хотите удалить урок «{sanitize_html(lesson.title)}»?",
        reply_markup=kb.as_markup()
    )

    await callback.answer()


# 2️⃣ Шаг: Фактическое удаление
@router.callback_query(F.data.startswith("lesson_delete_confirm:"))
async def lesson_delete_execute(callback: types.CallbackQuery):
    lesson_id = callback.data.split(":", 1)[1]
    lesson = await safe_get_lesson(lesson_id)
    if not lesson:
        await callback.answer("⚠️ Урок не найден", show_alert=True)
        return

    if not await sync_to_async(can_generate_lesson)(callback.from_user.id, lesson):
        await callback.answer("⚠️ Недостаточно прав", show_alert=True)
        return

    # Удаляем урок
    await sync_to_async(lesson.delete)()

    try:
        await callback.message.edit_reply_markup(reply_markup=None)
    except Exception:
        pass

    # Сообщаем пользователю
    await callback.message.answer(f"✅ Урок «{sanitize_html(lesson.title)}» успешно удалён!")

    # Возврат на главную или список уроков
    kb = InlineKeyboardBuilder()
    kb.button(text="🏠 На главную", callback_data="main_menu")
    kb.adjust(1)
    await callback.message.answer("Выберите действие:", reply_markup=kb.as_markup())

    await callback.answer()



# --- вспомогательные синхронные функции ---
def mark_lesson_discovered(lesson):
    lesson.is_discovered = True
    lesson.save(update_fields=["is_discovered"])

def mark_lesson_downloaded(lesson):
    lesson.is_downloaded = True
    lesson.save(update_fields=["is_downloaded"])

def get_blocks_list_sync(lesson):
    return list(lesson.blocks.order_by("order").all())

def build_docx_and_save(tmp_path, lesson_title, blocks):
    doc = Document()
    doc.add_heading(f"Урок: {lesson_title}", level=0)
    for block in blocks:
        doc.add_heading(f"{block.order}. {block.title}", level=1)
        html_content = markdown.markdown(block.content, extensions=["tables"])
        # Ваша синхронная sanitize_word должна уметь принимать doc
        sanitize_word(html_content, doc=doc)
    doc.save(tmp_path)
    return tmp_path

# Если track_user_activity — синхронная:
def track_user_activity_sync(user):
    metrics = track_user_activity(user)  # ваша синхронная функция
    metrics.increment_pdf_download()
    return True

# --- сам обработчик ---
@router.callback_query(F.data.startswith("lesson_download:"))
async def lesson_download(callback: types.CallbackQuery):
    try:
        await sync_to_async(check_rate_limit)(
            callback.from_user.id,
            "lesson_download",
            limit=5,
            window=60
        )
    except PermissionDenied as e:
        await callback.answer(f"⚠️ {str(e)}")
        return

    try:
        await get_or_create_user(callback.from_user.id, UserRole.SCHOOL_TEACHER, telegram_username=callback.from_user.username)
    except Exception as e:
        pass

    lesson_id = callback.data.split(":", 1)[1]
    lesson = await safe_get_lesson(lesson_id)
    if not lesson:
        await callback.answer("⚠️ Урок не найден", show_alert=True)
        return

    if not await sync_to_async(can_generate_lesson)(callback.from_user.id, lesson):
        await callback.answer("⚠️ Недостаточно прав", show_alert=True)
        return

    # Обновляем флаг discovered в треде
    try:
        await sync_to_async(mark_lesson_discovered)(lesson)
        await sync_to_async(mark_lesson_downloaded)(lesson)
    except Exception as e:
        print(e)
        pass

    # Получаем блоки (в треде)
    try:
        blocks = await sync_to_async(get_blocks_list_sync)(lesson)
    except Exception:
        await callback.answer("⚠️ Ошибка при получении блока(ов) урока", show_alert=True)
        return

    if not blocks:
        await callback.answer("⚠️ Урок пустой", show_alert=True)
        return

    # Создаём временный файл и генерируем docx в треде (чтобы не блокировать loop)
    tmp_path = None
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
            tmp_path = tmp.name

        # build_docx_and_save выполняется синхронно в отдельном потоке
        await sync_to_async(build_docx_and_save)(tmp_path, lesson.title, blocks)

        # Формируем клавиатуру
        kb = InlineKeyboardBuilder()
        kb.button(text="🔙 Об уроке", callback_data=f"lesson_detail:{lesson_id}")
        kb.button(text="🏠 На главную", callback_data="main_menu")
        kb.adjust(2)

        # Отправка документа (awaitable)
        await callback.message.answer_document(
            types.FSInputFile(tmp_path, filename=f"{lesson.title}.docx"),
            reply_markup=kb.as_markup()
        )

        await callback.answer("📄 Урок выгружен в Word")

    except Exception as e:
        # Ловим ошибки генерации/отправки
        kb = InlineKeyboardBuilder()
        kb.button(text="🏠 На главную", callback_data="main_menu")
        kb.adjust(1)
        await callback.message.answer(
            "❌ Не удалось выгрузить урок. Попробуйте снова.",
            reply_markup=kb.as_markup()
        )
        await callback.answer()
    finally:
        # Удаляем временный файл в треде (файловые операции также синхронные)
        if tmp_path and os.path.exists(tmp_path):
            try:
                await sync_to_async(os.remove)(tmp_path)
            except Exception:
                pass

    # Обновляем reply_markup - безопасно вызывать async метод
    try:
        await callback.message.edit_reply_markup(reply_markup=None)
    except Exception:
        pass

    # Обновляем метрики пользователя в треде (если sync)
    try:
        user = await get_user_by_tg(callback.from_user.id)
        # если track_user_activity — синхронная функция, используем wrapper
        await sync_to_async(track_user_activity_sync)(user)
    except Exception as e:
        # логируем, но не ломаем основной путь
        print(e)
        pass


@router.callback_query(F.data.startswith("lesson_actions"))
async def lesson_actions(callback: types.CallbackQuery):
    parts = callback.data.split(":")
    lesson_id = parts[1] if len(parts) > 1 else None
    block_index = int(parts[2]) if len(parts) > 2 else 1

    if not lesson_id:
        await callback.answer("⚠️ Неверные данные", show_alert=True)
        return

    lesson = await safe_get_lesson(lesson_id)
    if not lesson:
        await callback.answer("⚠️ Урок не найден", show_alert=True)
        return

    blocks = await sync_to_async(list)(lesson.blocks.order_by("order").all())
    total_blocks = len(blocks)
    if total_blocks == 0:
        await callback.answer("⚠️ Урок пустой", show_alert=True)
        return

    block_index = max(1, min(block_index, total_blocks))
    block = blocks[block_index - 1]

    text = (
        f"📘 <b>Урок:</b> {sanitize_html(lesson.title)}\n\n"
        f"<b>Блок {block.order}</b> — {sanitize_html(block.title)}\n\n"
        f"{sanitize_html(block.content)}\n\n"
        f"<i>Блок {block_index}/{total_blocks}</i>\n\n"
        f"<b>Выберите действие</b>"
    )

    kb = actions_kb_for_block(str(lesson.id), block_index)

    try:
        try:
            await callback.message.edit_reply_markup(reply_markup=None)
        except Exception:
            pass

        await safe_edit_text(
            callback.message,
            text,
            parse_mode="HTML",
            reply_markup=kb.as_markup()
        )
    except TelegramBadRequest:
        # если редактирование не получилось — просто отправим новое сообщение
        await callback.message.answer(text, parse_mode="HTML", reply_markup=kb.as_markup())

    await callback.answer()


# -----------------------
# Редактирование: показываем подсказку и переводим в состояние ожидания нового текста
# -----------------------
@router.callback_query(F.data.startswith("lesson_edit"))
async def lesson_edit_start(callback: types.CallbackQuery, state: FSMContext):
    parts = callback.data.split(":")
    lesson_id = parts[1] if len(parts) > 1 else None
    block_index = int(parts[2]) if len(parts) > 2 else 1

    lesson = await safe_get_lesson(lesson_id)
    if not lesson:
        await callback.answer("⚠️ Урок не найден", show_alert=True)
        return

    blocks = await sync_to_async(list)(lesson.blocks.order_by("order").all())
    if not blocks:
        await callback.answer("⚠️ Урок пустой", show_alert=True)
        return

    block_index = max(1, min(block_index, len(blocks)))
    block = blocks[block_index - 1]

    # Пояснение: нельзя предзаполнить поле ввода у пользователя через Telegram.
    # Поэтому показываем текущий текст и просим прислать новый.
    prompt = (
        f"✏️ <b>Редактирование блока {block.order} — {sanitize_html(block.title)}</b>\n\n"
        "Текущий текст блока:\n\n"
        f"<pre>{html.escape(block.content or '')}</pre>\n\n"
        "Отправьте новый текст, который заменит содержимое блока.\n\n"
    )

    kb = InlineKeyboardBuilder()
    kb.button(text="◀️ Отмена", callback_data=f"lesson_view:{lesson_id}:{block_index}")
    kb.adjust(1)

    try:
        try:
            await callback.message.edit_reply_markup(reply_markup=None)
        except Exception:
            pass

        await callback.message.answer(prompt, parse_mode="HTML", reply_markup=kb.as_markup())
    except TelegramBadRequest:
        await callback.answer("Не удалось открыть окно ввода.", show_alert=True)
        return

    # сохраняем в FSM данные для дальнейшей обработки
    await state.update_data(lesson_id=str(lesson.id), block_index=block_index, block_id=str(block.id))
    await state.set_state(EditBlockStates.waiting_for_new_content)

    await callback.answer()


# Принимаем новый текст от пользователя и сохраняем в блок
@router.message(EditBlockStates.waiting_for_new_content)
async def receive_new_block_content(message: types.Message, state: FSMContext):
    data = await state.get_data()
    lesson_id = data.get("lesson_id")
    block_id = data.get("block_id")
    block_index = data.get("block_index", 1)

    new_text = message.text.strip()
    if not new_text:
        await message.answer("Пустой текст — отправьте, пожалуйста, непустой текст или /cancel.")
        return

    # находим блок и сохраняем
    lesson = await safe_get_lesson(lesson_id)
    if not lesson:
        await message.answer("⚠️ Урок не найден. Операция прервана.")
        await state.clear()
        return

    if not await sync_to_async(can_generate_lesson)(message.from_user.id, lesson):
        await message.answer("⚠️ Недостаточно прав", show_alert=True)
        await state.clear()
        return

    try:
        block = await sync_to_async(lesson.blocks.get)(id=block_id)
    except Exception:
        await message.answer("⚠️ Блок не найден. Операция прервана.")
        await state.clear()
        return

    block.content = new_text
    await sync_to_async(block.save)()

    await state.clear()

    # Ответ и возврат в просмотр блока
    await message.answer("✅ Текст блока успешно обновлён.")
    # обновляем сообщение с блоком (если хотите — можно открыть lesson_view)
    await lesson_view_callback_simulate(message, lesson_id, block_index)

# Вспомогательная функция: обновление/переход к lesson_view (можно вызывать из разных мест)
async def lesson_view_callback_simulate(source_message, lesson_id: str, block_index: int):
    # reuse logic from lesson_view to display the block
    lesson = await safe_get_lesson(lesson_id)
    if not lesson:
        await source_message.answer("⚠️ Урок не найден")
        return
    blocks = await sync_to_async(list)(lesson.blocks.order_by("order").all())
    total_blocks = len(blocks)
    if total_blocks == 0:
        await source_message.answer("⚠️ Урок пустой")
        return
    block_index = max(1, min(block_index, total_blocks))
    block = blocks[block_index - 1]

    text = (
        f"📘 <b>Урок:</b> {sanitize_html(lesson.title)}\n\n"
        f"<b>Блок {block.order}</b> — {sanitize_html(block.title)}\n\n"
        f"{sanitize_html(block.content)}\n\n"
        f"<i>Блок {block_index}/{total_blocks}</i>"
    )
    kb = navigation_kb_for_block(str(lesson.id), block_index, total_blocks)
    try:
        await safe_edit_text(source_message, text, parse_mode="HTML", reply_markup=kb.as_markup())
    except TelegramBadRequest:
        await source_message.answer(text, parse_mode="HTML", reply_markup=kb.as_markup())

# -----------------------
# Удаление блока и пересчёт order
# -----------------------
@router.callback_query(F.data.startswith("lesson_delete_block"))
async def lesson_delete_block(callback: types.CallbackQuery):
    parts = callback.data.split(":")
    lesson_id = parts[1] if len(parts) > 1 else None
    block_index = int(parts[2]) if len(parts) > 2 else 1

    if not lesson_id:
        await callback.answer("⚠️ Неверные данные", show_alert=True); return

    lesson = await safe_get_lesson(lesson_id)
    if not lesson:
        await callback.answer("⚠️ Урок не найден", show_alert=True); return

    blocks = await sync_to_async(list)(lesson.blocks.order_by("order").all())
    total_blocks = len(blocks)
    if total_blocks == 0:
        await callback.answer("⚠️ Урок пустой", show_alert=True); return

    block_index = max(1, min(block_index, total_blocks))
    block_to_delete = blocks[block_index - 1]

    # удаляем блок
    await sync_to_async(block_to_delete.delete)()

    # пересчитываем order у оставшихся блоков
    remaining = await sync_to_async(list)(lesson.blocks.order_by("order").all())
    # гарантируем последовательность 1..N
    for i, b in enumerate(remaining, start=1):
        if b.order != i:
            b.order = i
            await sync_to_async(b.save)()

    # уведомляем пользователя и показываем следующий блок (если есть)
    new_total = len(remaining)
    if new_total == 0:
        await callback.message.answer("✅ Блок удалён. Урок теперь пустой.")
        # можно редиректнуть на детали урока
        try:
            await safe_edit_text(callback.message, f"📘 <b>Урок:</b> {sanitize_html(lesson.title)}\n\nУрок пуст.", parse_mode="HTML")
        except TelegramBadRequest:
            pass
        await callback.answer()
        return

    # если индекс был последний — показываем предыдущий
    new_index = min(block_index, new_total)
    try:
        await callback.message.edit_reply_markup(reply_markup=None)
    except Exception:
        pass

    await callback.answer("✅ Блок удалён")
    await lesson_view_callback_simulate(callback.message, str(lesson.id), new_index)


# -----------------------
# Улучшить ИИ: меню и действия
# -----------------------

def ai_menu_kb(lesson_id: str, block_index: int) -> InlineKeyboardBuilder:
    kb = InlineKeyboardBuilder()
    kb.button(text="🔼 Усложнить", callback_data=f"lesson_ai:{lesson_id}:{block_index}:complexify")
    kb.button(text="🔽 Упростить", callback_data=f"lesson_ai:{lesson_id}:{block_index}:simplify")
    kb.button(text="➕ Больше заданий", callback_data=f"lesson_ai:{lesson_id}:{block_index}:more_tasks")
    kb.button(text="➖ Убрать задания", callback_data=f"lesson_ai:{lesson_id}:{block_index}:remove_tasks")
    kb.button(text="◀️ Назад", callback_data=f"lesson_actions:{lesson_id}:{block_index}")
    kb.adjust(2)
    return kb


@router.callback_query(F.data.startswith("lesson_ai_menu"))
async def lesson_ai_menu(callback: types.CallbackQuery):
    parts = callback.data.split(":")
    lesson_id = parts[1]
    block_index = int(parts[2])

    try:
        await get_or_create_user(callback.from_user.id, UserRole.SCHOOL_TEACHER, telegram_username=callback.from_user.username)
    except Exception as e:
        pass

    lesson = await safe_get_lesson(lesson_id)
    if not lesson:
        await callback.answer("⚠️ Урок не найден", show_alert=True)
        return

    blocks = await sync_to_async(list)(lesson.blocks.order_by("order").all())
    block = blocks[max(0, min(block_index - 1, len(blocks) - 1))]

    try:
        await callback.message.edit_reply_markup(reply_markup=None)
    except Exception:
        pass

    text = (
        f"🤖 <b>Улучшение ИИ для блока {block.order} — {sanitize_html(block.title)}</b>\n\n"
        f"{sanitize_html(block.content)}\n\n"
        "<b>Выберите операцию</b>"
    )
    kb = ai_menu_kb(str(lesson.id), block_index)

    try:
        await safe_edit_text(callback.message, text, parse_mode="HTML", reply_markup=kb.as_markup())
    except TelegramBadRequest:
        await callback.message.answer(text, parse_mode="HTML", reply_markup=kb.as_markup())

    await callback.answer()


# Человеческие названия режимов для пользователя
MODE_NAMES_RU = {
    "complexify": "Усложнить",
    "simplify": "Упростить",
    "more_tasks": "Добавить задания",
    "remove_tasks": "Убрать задания"
}

# Основной обработчик AI-операций (запускает Celery-задачу)
@router.callback_query(F.data.startswith("lesson_ai:"))
async def lesson_ai_apply(callback: types.CallbackQuery):
    # формат: lesson_ai:{lesson_id}:{block_index}:{mode}
    parts = callback.data.split(":")
    if len(parts) < 4:
        await callback.answer("Неверные данные", show_alert=True)
        return

    try:
        await sync_to_async(check_rate_limit)(
            callback.from_user.id,
            "start_generation",
            limit=1,
            window=30
        )
    except PermissionDenied as e:
        await callback.answer(f"⚠️ {str(e)}")
        return

    lesson_id, block_index_s, mode = parts[1], parts[2], parts[3]
    block_index = int(block_index_s)

    lesson = await safe_get_lesson(lesson_id)
    if not lesson:
        await callback.answer("⚠️ Урок не найден", show_alert=True)
        return

    if not await sync_to_async(can_generate_lesson)(callback.from_user.id, lesson):
        await callback.answer("⚠️ Недостаточно прав", show_alert=True)
        return

    blocks = await sync_to_async(list)(lesson.blocks.order_by("order").all())
    if not blocks:
        await callback.answer("⚠️ Урок пустой", show_alert=True)
        return

    user = await get_user_by_tg(callback.from_user.id)

    # проверяем, остались ли генерации
    remaining = await sync_to_async(lambda: user.remaining_generations)()

    if remaining <= 0:
        kb = InlineKeyboardBuilder()
        kb.button(text="💎 Пополнить генерации", callback_data="shop")
        kb.button(text="🏠 На главную", callback_data="main_menu")
        kb.adjust(1)

        try:
            await callback.message.edit_reply_markup(reply_markup=None)
        except Exception:
            pass

        await safe_edit_text(
            callback.message,
            "⚠️ У вас закончились генерации.\n\n"
            "Пополните генерации <b>и продолжайте создавать задания!</b>",
            parse_mode="HTML",
            reply_markup=kb.as_markup()
        )
        await callback.answer()
        return

    # списываем одну генерацию
    await sync_to_async(user.decrement_generation)()

    block_index = max(1, min(block_index, len(blocks)))
    block = blocks[block_index - 1]

    # создаём ImproveStatus запись
    improve_status = await sync_to_async(ImproveStatus.objects.create)(
        block_id=block.id,
        mode=mode,
        status=ImproveStatus.Status.PENDING,
    )

    # запускаем Celery-задачу
    task = improve_block_task.delay(block.id, mode, improve_status.id)

    # сохраняем task_id
    improve_status.task_id = task.id
    await sync_to_async(improve_status.save)(update_fields=["task_id"])

    try:
        await callback.message.edit_reply_markup(reply_markup=None)
    except Exception:
        pass

    # получаем русское название режима для вывода пользователю
    mode_ru = MODE_NAMES_RU.get(mode, mode)

    await callback.message.answer(
        f"⚙️ Улучшаем блок ({mode_ru}).\n"
        f"Нажмите «🔄 Обновить статус», чтобы проверить готовность.",
        reply_markup=InlineKeyboardMarkup(
            inline_keyboard=[
                [InlineKeyboardButton(text="🔄 Обновить статус", callback_data=f"improve_status:{improve_status.id}")]
            ]
        ),
    )

    await callback.answer()


@router.callback_query(F.data.startswith("improve_status:"))
async def improve_status_handler(callback: types.CallbackQuery):
    _, improve_id = callback.data.split(":")

    try:
        await sync_to_async(check_rate_limit)(
            callback.from_user.id,
            "update_status",
            limit=10,
            window=30
        )
    except PermissionDenied as e:
        await callback.answer(f"⚠️ {str(e)}")
        return

    improve_status = await sync_to_async(
        ImproveStatus.objects.select_related("block__lesson").get
    )(id=int(improve_id))

    lesson_id = improve_status.block.lesson.id

    # Создаём клавиатуру
    kb = InlineKeyboardBuilder()
    # Кнопка "Обновить статус", если задача ещё выполняется
    if improve_status.status in [ImproveStatus.Status.PENDING, ImproveStatus.Status.IN_PROGRESS]:
        kb.button(text="🔄 Обновить статус", callback_data=f"improve_status:{improve_status.id}")
    kb.button(text="🔙 К уроку", callback_data=f"lesson_detail:{lesson_id}")
    kb.button(text="🏠 На главную", callback_data="main_menu")
    kb.adjust(1)

    if improve_status.status == ImproveStatus.Status.DONE:
        new_text = f"✅ Блок обновлён!\n\n{sanitize_html(improve_status.result_content)}"
        # редактируем текст и клавиатуру
        await callback.message.edit_text(new_text, parse_mode="HTML", reply_markup=kb.as_markup())

    elif improve_status.status == ImproveStatus.Status.FAILED:
        new_text = "❌ Ошибка при улучшении блока"
        await callback.message.edit_text(new_text, reply_markup=kb.as_markup())

    else:
        # PENDING или IN_PROGRESS — кнопки видны, текст не меняем, просто уведомляем
        await callback.answer("⚙️ Задача в процессе...", show_alert=False)
        try:
            # пробуем обновить клавиатуру на случай, если её ещё нет
            await callback.message.edit_reply_markup(reply_markup=kb.as_markup())
        except Exception:
            # игнорируем ошибку "message is not modified"
            pass







# -----------------------------
# Handlers: Main menu
# -----------------------------
# --- вспомогательные функции ---
@sync_to_async
def get_lessons_count(user_id: int) -> int:
    return Lesson.objects.filter(creator__telegram_id=user_id).count()


@sync_to_async
def get_lessons_page(user_id: int, offset: int, limit: int):
    return list(
        Lesson.objects.filter(creator__telegram_id=user_id)
        .order_by("-created_at")
        .all()[offset:offset+limit]
    )

@router.message(Command("main_menu"))
@router.callback_query(F.data == "main_menu")
async def main_menu(callback_or_message: types.Union[types.CallbackQuery, types.Message], state: FSMContext):
    try:
        await state.clear()
    except Exception:
        pass

    try:
        await get_or_create_user(callback_or_message.from_user.id, UserRole.SCHOOL_TEACHER, telegram_username=callback_or_message.from_user.username)
    except Exception as e:
        pass

    try:
        await sync_to_async(check_rate_limit)(
            callback_or_message.from_user.id,
            "main_menu",
            limit=1,
            window=5
        )
    except PermissionDenied as e:
        return

    user_id = callback_or_message.from_user.id

    user = await get_user_by_tg(callback_or_message.from_user.id)
    metrics = await sync_to_async(track_user_activity)(user)

    total = await get_lessons_count(user_id)
    remaining_generations = user.remaining_generations
    text = f"📚 Создано материалов: <b>{total}</b>\n"
    text += f"🚀 Осталось генераций: <b>{remaining_generations}</b>"

    kb = InlineKeyboardBuilder()
    kb.button(text="Создать урок", callback_data="create_lesson")
    kb.button(text="Материалы", callback_data="my_lessons:0")
    kb.button(text="Тариф", callback_data="shop")
    kb.button(text="Настройки", callback_data="settings")
    kb.button(text="Помощь", callback_data="help")
    kb.adjust(2)

    if isinstance(callback_or_message, types.Message):
        await callback_or_message.answer(
            text,
            reply_markup=kb.as_markup(),
            parse_mode="HTML"
        )
    else:
        # Скрываем кнопки предыдущего сообщения
        if callback_or_message.message.reply_markup:
            try:
                await callback_or_message.message.edit_reply_markup(reply_markup=None)
            except Exception:
                pass

        # Отправляем новое сообщение в чат
        await callback_or_message.message.answer(
            text,
            reply_markup=kb.as_markup(),
            parse_mode="HTML"
        )

        # Закрываем "Loading…" у callback
        await callback_or_message.answer()

@router.callback_query(F.data == "create_lesson")
async def create_lesson_start(callback: types.CallbackQuery, state: FSMContext):
    await state.set_state(CreateLessonStates.waiting_for_title)

    try:
        await get_or_create_user(callback.from_user.id, UserRole.SCHOOL_TEACHER, telegram_username=callback.from_user.username)
    except Exception as e:
        pass

    try:
        await callback.message.edit_reply_markup(reply_markup=None)
    except Exception:
        pass

    await callback.message.answer(
        "Введите тему нового урока:"
    )

    await callback.answer()




# -----------------------------
# Handlers: My lessons (list with pagination)
# -----------------------------
@router.callback_query(F.data.startswith("my_lessons:"))
async def my_lessons(callback: types.CallbackQuery):
    try:
        await get_or_create_user(callback.from_user.id, UserRole.SCHOOL_TEACHER, telegram_username=callback.from_user.username)
    except Exception as e:
        pass

    page = int(callback.data.split(":", 1)[1])
    per_page = 5
    offset = page * per_page

    lessons = await get_lessons_page(callback.from_user.id, offset, per_page)
    total = await get_lessons_count(callback.from_user.id)

    if not lessons:
        await callback.answer("Уроков нет", show_alert=True)
        return

    text = "📖 <b>Мои материалы:</b>\n\n"
    kb = InlineKeyboardBuilder()

    for lesson in lessons:
        text += f"• {sanitize_html(lesson.title)}\n"
        kb.button(text=lesson.title, callback_data=f"lesson_detail:{lesson.id}")

    # Навигация
    if offset > 0:
        kb.button(text="⬅️ Назад", callback_data=f"my_lessons:{page-1}")
    if offset + per_page < total:
        kb.button(text="➡️ Далее", callback_data=f"my_lessons:{page+1}")

    kb.button(text="🏠 На главную", callback_data="main_menu")
    kb.adjust(1)

    try:
        await callback.message.edit_reply_markup(reply_markup=None)
    except Exception:
        pass

    await safe_edit_text(
        callback.message,
        text,
        parse_mode="HTML",
        reply_markup=kb.as_markup()
    )

    await callback.answer()


# -----------------------------
# Handlers: Settings
# -----------------------------
@router.callback_query(F.data == "settings")
async def settings(callback: types.CallbackQuery):
    kb = InlineKeyboardBuilder()
    kb.button(text="🏠 На главную", callback_data="main_menu")
    kb.adjust(1)

    try:
        await callback.message.edit_reply_markup(reply_markup=None)
    except Exception:
        pass

    # Отправляем новое сообщение в чат (не редактируем старое)
    await callback.message.answer(
        "⚙️ Чтобы изменить настройки, отправьте /start",
        reply_markup=kb.as_markup()
    )

    # Закрываем "Loading…" у callback
    await callback.answer()

@router.callback_query(F.data == "help")
async def help_cmd(callback: types.CallbackQuery, state: FSMContext):
    kb = InlineKeyboardBuilder()
    kb.button(text="🏠 На главную", callback_data="main_menu")
    kb.adjust(1)

    user_id = callback.from_user.id

    try:
        try:
            await callback.message.edit_reply_markup(reply_markup=None)
        except Exception:
            pass

        await callback.message.answer(
            "ℹ️ Напишите сообщение в тех. поддержку\n"
            "Мы отвечаем в течение часа.\n",
            reply_markup=kb.as_markup()
        )
    except Exception:
        pass

    await state.set_state(SupportStates.waiting_message)

    await callback.answer()

# ----- Состояния техподдержки -----
class SupportStates(StatesGroup):
    waiting_message = State()

@router.message(SupportStates.waiting_message)
async def support_message_handler(message: types.Message):
    """Сохраняем сообщение пользователя в тикет и обновляем статус на RECEIVED"""
    user_id = message.from_user.id

    username = message.from_user.username

    # проверка лимита сообщений (10/час)
    hour_ago = timezone.now() - timedelta(hours=1)

    recent_count = await sync_to_async(lambda: TicketMessage.objects.filter(
        ticket__user_id=user_id,
        created_at__gte=hour_ago
    ).count())()

    if recent_count >= 10:
        await message.answer("⚠️ Вы превысили лимит: максимум 10 сообщений в час.")
        return

    # Получаем существующий тикет или создаём новый
    ticket, created = await sync_to_async(SupportTicket.objects.get_or_create)(
        user_id=user_id,
        defaults={
            "username": username,
            "ticket_id": f"T-{timezone.now().strftime('%Y%m%d%H%M%S')}",
            "status": SupportTicket.Status.RECEIVED,
        }
    )

    # Если тикет уже существовал, обновляем его статус на RECEIVED
    if not created:
        ticket.status = SupportTicket.Status.RECEIVED
        await sync_to_async(ticket.save)()


    # Определяем вложение
    # Начальные значения
    text = message.text or None
    attachment_id = None

    # Фото
    if message.photo:
        attachment_id = message.photo[-1].file_id  # лучше брать последнее фото, оно с наибольшим разрешением
        if message.caption:
            text = message.caption

    # Документ
    elif message.document:
        attachment_id = message.document.file_id
        if message.caption:
            text = message.caption

    # Голосовое сообщение
    elif message.voice:
        attachment_id = message.voice.file_id
        if message.caption:
            text = message.caption

    # Создаём сообщение
    await sync_to_async(TicketMessage.objects.create)(
        ticket=ticket,
        text=text,
        attachment_id=attachment_id
    )

    try:
        await message.edit_reply_markup(reply_markup=None)
    except Exception:
        pass

    # Формируем клавиатуру с кнопкой "На главную"
    keyboard = InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text="🏠 На главную", callback_data="main_menu")]
    ])

    await message.answer(
        "✅ Сообщение передано в поддержку.",
        reply_markup=keyboard
    )
